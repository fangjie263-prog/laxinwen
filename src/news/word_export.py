"""Word（DOCX）研究阅读包导出。

把与 Portable HTML 完全同一批新闻数据（``Storage.list_articles_with_analysis`` +
``Storage.is_usable``，排序一致：按发布时间倒序）生成一个适合研究阅读的
``.docx`` 简报：

- **目录**：首页给出「标题 → 正文」的**内部超链接**（bookmark + hyperlink），
  打开 Word 后**点击目录条目即可直接跳到对应文章**，无需手动“更新域”。
- **双向导航**：每篇新闻正文均含「↑ 返回目录」内部超链接，指向目录的
  ``laxinwen-toc`` bookmark，实现「目录 → 正文 → 返回目录 → 下一篇」的完整闭环。
- **每篇新闻**：完整保留标题、来源、发布时间（北京时间）、正文，以及
  **可点击的原文 URL**（真正写入 Word 的 HYPERLINK relationship）。
- **纯 Python 生成标准 .docx**：仅依赖 ``python-docx``，不需要
  LibreOffice / Word COM / Windows Office。

设计约束（与 Portable 导出保持一致）：
- **不触碰** pipeline / discovery / fetch / extraction / storage / scheduler；
  Word 只是 export 层；
- **复用** 现有数据读取与排序（``list_articles_with_analysis`` 发布时间倒序）；
- **不重新抓网页**：只使用数据库中已存储的字段；
- 展示时间统一北京时间（Asia/Shanghai），24 小时制（复用 ``beijing``）；
- ``usable=0`` 不影响任务状态：导出仍可正常生成（空包 / 目录为空），
  由调用方（scheduled-fetch）决定 FETCH/EXPORT 标记。
"""

from __future__ import annotations

import logging
import sqlite3
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor
except Exception:  # pragma: no cover - 依赖缺失时给出清晰报错
    Document = None  # type: ignore
    WD_ALIGN_PARAGRAPH = None  # type: ignore
    OxmlElement = None  # type: ignore
    qn = None  # type: ignore
    Pt = None  # type: ignore
    RGBColor = None  # type: ignore

from .beijing import fmt_date as _bj_fmt_date, fmt_dt as _bj_fmt_dt
from .news_archive import _parse_datetime
from .storage import Storage

logger = logging.getLogger(__name__)


# 默认导出根目录（与 GUI / CLI / scheduled-fetch 保持一致）
DEFAULT_WORD_DIR = Path("data") / "export" / "word"

_ACCENT = RGBColor(0x6B, 0x4F, 0x2A)   # 阅读器暖棕强调色
_SOFT = RGBColor(0x5C, 0x5C, 0x58)     # 次级文本
_MUTED = RGBColor(0x8A, 0x8A, 0x84)    # 弱文本


@dataclass
class WordExportResult:
    """Word 导出统计（字段与 PortableResult 保持一致，便于统一消费）。"""

    exported: int = 0
    failed: int = 0
    analyzed_ok: int = 0
    analyzed_failed: int = 0
    unanalyzed: int = 0
    files: list[Path] = field(default_factory=list)
    index_path: Optional[Path] = None

    def as_dict(self) -> dict:
        return {
            "exported": self.exported,
            "failed": self.failed,
            "analyzed_ok": self.analyzed_ok,
            "analyzed_failed": self.analyzed_failed,
            "unanalyzed": self.unanalyzed,
        }


# --------------------------------------------------------------------------
# 低层 OOXML 辅助：bookmark / internal hyperlink / external hyperlink
# --------------------------------------------------------------------------

def _add_bookmark(paragraph, bookmark_name: str) -> None:
    """给段落开头插入一个命名 bookmark（供目录内部超链接跳转）。

    Word 内部跳转 = ``bookmarkStart/End`` 定义锚点 + 目录里的
    ``w:hyperlink w:anchor=bookmark``。用户点击目录条目即跳转到该 bookmark。
    """
    if OxmlElement is None:
        return
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(_new_bookmark_id()))
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), start.get(qn("w:id")))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


_bookmark_counter = 1000


def _new_bookmark_id() -> int:
    global _bookmark_counter
    _bookmark_counter += 1
    return _bookmark_counter


def _add_internal_hyperlink(paragraph, anchor: str, text: str, *, bold: bool = False) -> None:
    """给段落添加一个**内部**超链接（跳转到文档内的 bookmark ``anchor``）。

    这是实现「目录 → 新闻正文」点击跳转的关键。
    """
    if OxmlElement is None:
        paragraph.add_run(text)
        return
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")  # Word 标准超链接蓝
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _add_external_hyperlink(paragraph, url: str, text: str) -> None:
    """给段落添加一个**外部**超链接（跳转到原始网页 URL）。

    会写入 ``document.xml.rels`` 中的 External Relationship，
    因此 Word 里点击即可在浏览器打开原文。
    """
    if OxmlElement is None:
        paragraph.add_run(text or url)
        return
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# --------------------------------------------------------------------------
# 文档渲染
# --------------------------------------------------------------------------

def _set_style_font(doc: Document, name: str = "等线") -> None:
    """设置文档默认字体（兼容中文）。"""
    if qn is None:
        return
    style = doc.styles["Normal"]
    style.font.name = name
    style.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _add_title_line(doc: Document, text: str, size: int = 22, color=None, bold: bool = True) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _add_meta_line(doc: Document, text: str, size: int = 12, color=None) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _normalize_rows(rows) -> list[dict]:
    """把 SQLite 行/字典统一转成 dict 列表。"""
    return [dict(r) for r in rows]


# 目录锚点 bookmark 名：每篇新闻的「返回目录」内部超链接都指向这里。
# 所有 Word 文档统一使用该名称，确保「正文 → 目录」导航稳定可解析。
_TOC_BOOKMARK = "laxinwen-toc"


def _bookmark_name(index: int, article_id: int) -> str:
    """为每篇新闻生成稳定、唯一的 bookmark 名。"""
    return f"art-{index}-{article_id}"


def render_word_docx(
    rows,
    *,
    source_name: str,
    source_id: str,
    job_id: str = "",
    total: int,
    generated_at: datetime,
    body_html: bool = False,
) -> Document:
    """渲染一份完整的研究阅读包 ``.docx``（返回 ``docx.Document``）。

    ``rows`` 与 portable 完全同一批（已 ``is_usable`` 过滤、已按发布时间倒序）。

    双向导航：
    - 顶部目录含 ``laxinwen-toc`` bookmark，目录条目用内部超链接跳转到正文；
    - 每篇新闻正文均含「↑ 返回目录」内部超链接，指向 ``laxinwen-toc``，
      从而实现「目录 → 正文 → 返回目录」的完整双向导航。
    """
    if Document is None:  # pragma: no cover
        raise RuntimeError("python-docx 未安装，无法生成 Word 导出。请先安装依赖。")

    rows = _normalize_rows(rows)
    doc = Document()
    _set_style_font(doc)

    # ---- 封面区 ----
    _add_title_line(doc, "Laxinwen 新闻阅读包", size=26, color=_ACCENT)
    today = _bj_fmt_date(generated_at)
    _add_meta_line(doc, today, size=14, color=_SOFT)
    job_label = f"{source_name.upper()} · {job_id}" if job_id else source_name.upper()
    _add_meta_line(doc, job_label, size=14, color=_SOFT)
    _add_meta_line(doc, f"共 {total} 篇", size=13, color=_MUTED)
    _add_meta_line(
        doc,
        "由 laxinwen 生成 · 点击目录条目可跳转到对应新闻 · 原文链接可在浏览器打开",
        size=10,
        color=_MUTED,
    )
    doc.add_paragraph()

    # ---- 目录（内部超链接） ----
    toc_title = doc.add_paragraph()
    trun = toc_title.add_run("目录")
    trun.font.size = Pt(18)
    trun.bold = True
    trun.font.color.rgb = _ACCENT
    # 目录锚点：每篇正文的「返回目录」内部超链接跳转到此处。
    _add_bookmark(toc_title, _TOC_BOOKMARK)

    for i, row in enumerate(rows, start=1):
        title = row.get("title") or f"（无标题）#{row.get('id')}"
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(12)
        _add_internal_hyperlink(p, _bookmark_name(i, int(row["id"])), f"{i}. {title}")

    # 分页，让正文从新页开始
    doc.add_page_break()

    # ---- 正文（每篇一个 bookmark 锚点） ----
    for i, row in enumerate(rows, start=1):
        article_id = int(row["id"])
        published = (
            _parse_datetime(row.get("published_at"))
            or _parse_datetime(row.get("discovered_at"))
        )

        # 标题（Heading 1 + bookmark）
        heading = doc.add_heading(level=1)
        heading.add_run(f"{i}. {row.get('title') or '（无标题）'}")
        _add_bookmark(heading, _bookmark_name(i, article_id))

        # 来源 / 发布时间
        src = row.get("source_name") or row.get("source_id") or source_name
        meta = f"来源：{src}    发布时间：{_bj_fmt_dt(published) if published else '—'}"
        mp = doc.add_paragraph()
        mrun = mp.add_run(meta)
        mrun.font.size = Pt(10)
        mrun.font.color.rgb = _SOFT

        # 原文 URL（可点击的外部超链接）
        url = row.get("canonical_url") or ""
        if url:
            up = doc.add_paragraph()
            label = up.add_run("原文链接：")
            label.font.size = Pt(10)
            label.font.color.rgb = _SOFT
            _add_external_hyperlink(up, url, url)

        # 「返回目录」：真正的 Word 内部超链接，跳转到顶部目录 bookmark。
        # 每篇新闻各自独立拥有，便于从任意一篇直接回到目录。
        back = doc.add_paragraph()
        back.paragraph_format.left_indent = Pt(0)
        _add_internal_hyperlink(back, _TOC_BOOKMARK, "↑ 返回目录")

        # 正文
        body = (row.get("body_text") or "").strip()
        if body:
            for para in body.split("\n"):
                para = para.strip()
                if not para:
                    continue
                bp = doc.add_paragraph()
                bp.paragraph_format.first_line_indent = Pt(24)
                bp.paragraph_format.line_spacing = 1.5
                bp.add_run(para)
        else:
            empty = doc.add_paragraph()
            erun = empty.add_run("（本篇无正文）")
            erun.font.color.rgb = _MUTED
            erun.font.size = Pt(11)

        # 每篇文章之间加分隔
        doc.add_paragraph()

    return doc


# --------------------------------------------------------------------------
# 导出入口
# --------------------------------------------------------------------------

def default_word_path(source_id: str, job_id: str = "") -> Path:
    """默认 Word 输出路径：``data/export/word/Laxinwen-<SOURCE>-<日期>-<时间>-<job>.docx``。"""
    job_suffix = f"-{job_id}" if job_id else ""
    return (
        DEFAULT_WORD_DIR
        / (
            f"Laxinwen-{source_id.upper()}-"
            f"{datetime.now().strftime('%Y-%m-%d-%H%M%S')}{job_suffix}.docx"
        )
    )


def _collect_stats(rows) -> tuple[int, int, int]:
    """统计 analyzed_ok / analyzed_failed / unanalyzed（与 portable 语义一致）。"""
    analyzed_ok = analyzed_failed = unanalyzed = 0
    for row in rows:
        status = _ai_status(row)
        if status == "ok":
            analyzed_ok += 1
        elif status == "failed":
            analyzed_failed += 1
        else:
            unanalyzed += 1
    return analyzed_ok, analyzed_failed, unanalyzed


def _ai_status(row) -> str:
    """返回文章 AI 状态：'ok' / 'failed' / 'none'（复用 news_archive 语义）。"""
    if isinstance(row, sqlite3.Row):
        row = dict(row)
    if row.get("ai_status"):
        return "ok"
    if row.get("ai_has_failed"):
        return "failed"
    return "none"


def export_word_package(
    storage: Storage,
    out_path: str | Path,
    *,
    source_id: Optional[str] = None,
    limit: int = 100,
    job_id: str = "",
) -> WordExportResult:
    """导出 Word 研究阅读包：``out_path``（单个 ``.docx`` 文件）。

    与 portable 使用**同一批新闻数据**（``list_articles_with_analysis`` +
    ``is_usable``，按发布时间倒序）。不同 job 通过文件名中的 job id 区分，
    互不覆盖。
    """
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    rows = storage.list_articles_with_analysis(
        source_id=source_id, limit=limit if limit else 10**9
    )
    rows = [r for r in rows if storage.is_usable(r)]
    result = WordExportResult()
    result.exported = len(rows)
    result.analyzed_ok, result.analyzed_failed, result.unanalyzed = _collect_stats(rows)

    source_name = rows[0]["source_name"] if rows else source_id
    doc = render_word_docx(
        rows,
        source_name=source_name,
        source_id=source_id or "",
        job_id=job_id,
        total=len(rows),
        generated_at=datetime.now(timezone.utc),
    )
    doc.save(str(out_path))
    result.index_path = out_path
    result.files.append(out_path)
    logger.info(
        "Word 导出完成: %d 篇 → %s（已分析 %d / 失败 %d / 未分析 %d）",
        result.exported,
        out_path,
        result.analyzed_ok,
        result.analyzed_failed,
        result.unanalyzed,
    )
    return result
